from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from docx.shared import Pt, Cm
from docx import Document
import pyautogui as pyg
import win32com.client
import pandas as pd
import pythoncom
import warnings
import textwrap
from time import sleep
import os
import sys


# Configuration
warnings.filterwarnings("ignore", category=UserWarning, module="openpyxl")
excel_path = r'path\to\project\data\WorkOrders.xlsm'
service_order_txt_path = r'path\to\project\scripts\service_order_counter.txt'

# Read current service orders from Excel and tracking file
df = pd.read_excel(excel_path, sheet_name='Records')
df = df.set_index('WO')
excel_wo = max(df.index)

with open(service_order_txt_path, 'r') as file:
    txt_wo = int(file.read())

print(f'WO in file: {txt_wo} | WO in Excel: {excel_wo}')

# Exit if no new service orders
if txt_wo >= excel_wo:
    print('No new service order requests')
    sleep(1)
    sys.exit()

def add_service_order_number(row):
    """Add service order number to document header"""
    print('service order:', row.name)
    cell = table0.cell(0, 2)
    paragraph = cell.paragraphs[0]
    paragraph.add_run(str(row.name))
    run = paragraph.runs[0]
    run.font.size = Pt(22)
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def write_service_order_data(row):
    """Populate service order form with data from Excel"""
    table_positions = {
        'NAME': [0, 3],
        'DEPARTMENT': [0, 5],
        'COST_CENTER': [0, 11],
        'DATE': [1, 9],
        'DEADLINE': [1, 12]
    }

    # Fill in basic information
    for col, value in row.items():
        if col in table_positions:
            cell = table1.cell(table_positions[col][0], table_positions[col][1])
            cell.text = str(value)
            cell.paragraphs[0].runs[0].font.size = Pt(14)
            print(f'{col}: {value}')

    # Handle service description with text wrapping
    description = row['SERVICE_DESCRIPTION']
    print('DESCRIPTION:', description)

    line_length = 105
    lines = textwrap.wrap(description, width=line_length, break_long_words=False)
    
    if len(lines) > 4:
        print('DESCRIPTION TOO LONG (>4 lines)')
        sleep(3)
        sys.exit()

    # Add description lines to document
    for i, line in enumerate(lines, 2):
        cell = table1.cell(i, 1)
        cell.text = str(line).upper()
        cell.paragraphs[0].runs[0].font.size = Pt(14)

def insert_signatures(row):
    """Insert digital signatures for requestor and supervisor"""
    print('--- Inserting signatures ---')
    
    # Load signature mappings from backend sheet
    df_backend = pd.read_excel(excel_path, sheet_name='Backend')
    df_backend = df_backend.drop(columns=df_backend.columns[2:])

    df_names = pd.DataFrame()

    # Find supervisor for current requestor
    for j in df_backend.index:
        if row['NAME'] == df_backend.at[j, 'REQUESTOR']:
            print("Requestor:", row['NAME'], "| Supervisor:", df_backend.at[j, 'SUPERVISOR'])
            data = {
                'NAME': [row['NAME'], df_backend.at[j, 'SUPERVISOR']],
                'POSITION': [[6, 2], [6, 8]],
                'ROLE': ['REQUESTOR', 'SUPERVISOR']
            }
            df_names = pd.DataFrame(data)
            break

    if df_names.empty:
        print("Warning: Request from someone without signature mapping, verify if WO is valid.")
        sleep(3)
        return

    # Insert signature images
    signatures_path = r'path\to\project\signatures'
    for i in df_names.index:
        pos = df_names.at[i, 'POSITION']
        cell = table1.cell(pos[0], pos[1])
        cell.paragraphs[0].alignment = 1  # Center alignment
        img_path = os.path.join(signatures_path, df_names.at[i, 'NAME'] + '.PNG')
        cell.paragraphs[0].add_run().add_picture(img_path, height=Cm(1.3))

def print_word_document(printer_names: list, word_file_path: str):
    """Print Word document to specified printers"""
    word_instance = None
    print_doc = None
    
    try:
        pythoncom.CoInitialize()
        filename = os.path.basename(word_file_path)

        # Check if Word is already running
        try:
            word_instance = win32com.client.GetActiveObject("Word.Application")
            word_was_not_running = False
        except:
            word_instance = win32com.client.Dispatch("Word.Application")
            word_was_not_running = True

        print_doc = word_instance.Documents.Open(word_file_path)

        # Print to each specified printer
        for printer in printer_names:
            print(f'Printing {filename} to {printer}')
            word_instance.ActivePrinter = printer
            print_doc.PrintOut()
            sleep(1)

    except Exception as e:
        # Log printing errors
        with open(r'path\to\project\logs\printer_backlog.txt', 'a') as backlog:
            error_msg = f'\n{datetime.now().strftime("%Y-%m-%d %H:%M:%S")} - Error printing {filename} to {printer}:\n{e}\n\n'
            print(error_msg)
            backlog.write(error_msg)

    finally:
        if print_doc is not None:
            print_doc.Close(False)
        if word_instance is not None and word_was_not_running:
            word_instance.Quit()
        pythoncom.CoUninitialize()

# Main processing loop
print('')
for service_number in range(txt_wo + 1, excel_wo + 1):
    
    # Load Word template
    doc = Document(r"path\to\project\templates\service_order_template.docx")
    table0 = doc.tables[0]
    table1 = doc.tables[1]

    # Get service order data
    service_order = df.loc[service_number].copy()
    service_order = service_order.drop(service_order.index[-3:], axis=0)
    
    # Validate data completeness
    if any(service_order.isna()):
        print(f'ERROR: service order {service_number} has empty values, please correct in Excel.')
        sleep(3)
        sys.exit()
    
    # Format dates and description
    service_order['DATE'] = service_order['DATE']
    service_order['DEADLINE'] = service_order['DEADLINE']
    service_order['SERVICE_DESCRIPTION'] = str(service_order['SERVICE_DESCRIPTION']).upper()

    # Generate document
    add_service_order_number(service_order)
    write_service_order_data(service_order)
    print('')
    insert_signatures(service_order)
    print('')

    # Save generated document
    word_file_path = os.path.join(r'path\to\project\output', f"WO_{service_number}.docx")
    doc.save(word_file_path)
    sleep(2)

    # Print to designated printers
    printers = [
        r'\\print-server\department-printer',
        r'\\print-server\maintenance-printer',
        r'\\print-server\color-printer'
    ]
    print_word_document(printers, word_file_path)
    sleep(2)

    # Update tracking file
    with open(service_order_txt_path, 'w') as file:
        file.write(str(service_number))
        print(f'\nUpdated tracking file: {service_number}')